#!/usr/bin/env python
# coding: utf-8

# In[21]:


#pip install python-docx


# In[8]:


import numpy as np
import os
import nltk
from nltk.stem import WordNetLemmatizer
from nltk.corpus import wordnet
from nltk.stem import PorterStemmer
from nltk.tokenize import TweetTokenizer
from nltk.tokenize import word_tokenize
from natsort import natsorted
from nltk.corpus import stopwords
import string


# In[12]:


def read_file(filename):
    with open(filename, 'r', encoding ="ascii", errors ="surrogateescape") as f:
        stuff = f.read()
    f.close()
    return stuff
  
def preprocessing(final_string):
    # Tokenize.
    only_words = word_tokenize(final_string)
  
    # Lemmatization
    wordnet_lemmatizer = WordNetLemmatizer()
    afterLemitKeywords=[w for w in only_words if w in wordnet_lemmatizer.lemmatize(w)]

    # List of English stop words 
    english_stop_words=set(stopwords.words("english"))
    # Removal of stop words from the text
    token_list=[w for w in afterLemitKeywords if not w in english_stop_words]
    return token_list
  
# In this example, we create the positional index for only 1 folder.
folder_names = ["Assignment"]
    
# Initialize the file no.
fileno = 0
  
# Initialize the dictionary.
pos_index = {}
  
# Initialize the file mapping (fileno -> file name).
file_map = {}
  

for folder_name in folder_names:
  
    # Open files.
    file_names = natsorted(os.listdir(folder_name))
  
    # For every file.
    for file_name in file_names:
        
        # Read file contents.
        stuff = read_file(folder_name + "/" + file_name)

        final_token_list = preprocessing(stuff)

        import json
        # open output file for writing
        with open(folder_name + "/" + "preprocessed_"+file_name, "w") as filehandle:
            json.dump(final_token_list, filehandle)
   
        # For position and term in the tokens.
        for pos, term in enumerate(final_token_list):

                    # If term already exists in the positional index dictionary.
                    if term in pos_index:
                          
                        # Increment total freq by 1.
                        pos_index[term][0] = pos_index[term][0] + 1
                          
                        # Check if the term has existed in that DocID before.
                        if fileno in pos_index[term][1]:
                            pos_index[term][1][fileno].append(pos)
                              
                        else:
                            pos_index[term][1][fileno] = [pos]
  
                    # If term does not exist in the positional index dictionary 
                    # (first encounter).
                    else:
                          
                        # Initialize the list.
                        pos_index[term] = []
                        pos_index[term].append(1)
                        pos_index[term].append({})      
                        # Add doc ID to postings list.
                        pos_index[term][1][fileno] = [pos]

        with open(folder_name + "/" + "inverted_index_"+file_name, "w") as file:
            file.write(json.dumps(pos_index))
        print(pos_index)
        # Map the file no. to the file name.
        file_map[fileno] = folder_name + "/" + file_name
              
        fileno += 1
  
# Sample positional index to test the code.
sample_pos_idx = pos_index["SAP"]
print("Positional Index")
print(sample_pos_idx)
  
file_list = sample_pos_idx[1]
print("Filename, [Positions]")
for fileno, positions in file_list.items():
    print(file_map[fileno], positions)


# # Boolean Query

# In[20]:


from nltk.stem import PorterStemmer
from nltk.tokenize import sent_tokenize, word_tokenize
from docx import Document
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
import nltk
import glob
import errno
import collections
import os
#initializing data structures
dictG=[['',0]] #Final Dictionary
queryG=[] #Query
listC=[['',0]] 
docsG=[] #Docslist

#Proccess docs for index building
def buildIndex(fileName, docID):

    words = set() #set to store words from each doc
    
    #from each doc, append words into raw words list
    document=Document(fileName)
    for p in document.paragraphs:
        listT=p.text.split()
        for word in listT:
            words.add(word)    
    #print(words)
    
    #Removing Stop words
    stop_words=list(stopwords.words('english'))    
    filteredQ=[w for w in list(words) if not w in stop_words]
    #print('***************** filtered sentence after removing stop words*****************')
    #print(filteredQ)
    
    #Lemmentizing
    lemmatizer = WordNetLemmatizer()
    lemmaList=[]
    for word in filteredQ:
        lemmaList.append(lemmatizer.lemmatize(word))
    #print('*************** After lemmatization**********')
    #print(lemmaList)
    
    flagD=0
    flagDuplicateDoc=0
    for word in lemmaList:
        for element in dictG:
            if(element[0]==word):
                flagD=1
                for x in element[1:]:
                    if(x==docID):
                        flagDuplicateDoc=1
                        break
                if(flagDuplicateDoc==0):
                    element.append(docID)
                if(flagDuplicateDoc==1):
                    flagDuplicateDoc=0
        if(flagD==0):
            dictG.append([word,docID])
        if(flagD==1):
            flagD=0

'''Proccesses input Boolean query'''
def processQuery():
    global queryG
    flagQ=0
    queryRaw=input('Enter your query - ')
    queryT=queryRaw.split()
    #lemmeantize the query
    lemmatizer = WordNetLemmatizer()
    for word in queryT:
        queryG.append(lemmatizer.lemmatize(word))
    
    obtainTermsFromDictionary()
    getRequiredDocs()

''' Obtains terms from Index - Returns - list of 
terms+docID ['term',id,id...]'''
def obtainTermsFromDictionary():
    global dictG
    global listC
    listC=[['',0]]
    for word in queryG:
        for element in dictG:
            if(word==element[0]):
                listC.append(element)
    listC.pop(0)

'''Gets the associated doc IDs, by infix evaluation of list of doc Ids using stack'''
def getRequiredDocs():
    
    print(queryG) #given the current query
    ListA = list() #temporary list of docs for boolean query
    termCount = 0 #current termcount
    templist = listC[termCount][1:] #doc list from 1st term    
    stack = list(); #stack to manage the query
    
    while len(queryG) > 0:
        
        c = queryG.pop(0)

        if isinstance(c, collections.Sequence):
            if c in ['and','or','not']: 
                stack.append(c)
            else:
                for e in listC:
                    if e[0] == c:
                        stack.append(listC[listC.index(e)][1:])
        
        if ')' in c:
            
            num2 = stack.pop()
            op = stack.pop()
            num1 = stack.pop()
            
            if op == "and":
                stack.append(And(num1,num2))
            if op == "or":
                stack.append(Or(num1,num2))
            if op == "not":
                stack.append(Not(num1,num2))
          
    print('The Docs against the query are - ')
    print(stack[0])

'''Logical Not operation - Args(list,list)
returns - list'''
def Not(list1,list2):
    
    ListA = list()
    if(len(list1) >= len(list2)):
        for x in list1:
            if x not in list2:
                ListA.append(x)
    elif(len(list2) >= len(list1)):            
        for x in list2:
             if x not in list1:
                ListA.append(x)
    return ListA

'''Logical Or operation - Args(list,list)
returns - list'''
def Or(list1,list2):
    
    ListA = list()
    for x in list1:
        ListA.append(x)
        for y in list2:
            if(x != y):
                ListA.append(y)
    
    return list(set(ListA))

'''Logical And operation - Args(list,list)
returns - list'''
def And(list1,list2):
    
    ListA = list()
    for x in list1:
        for y in list2:
            if(x == y):
                ListA.append(x)
    
    return ListA

'''Final running of code'''
#Replace path according to your working directory structure

path = os.path.join(os.path.dirname('C:\\Users\\priyesh-s\\Desktop\\IT Assignment\\'), '*.docx')
files=glob.glob(path)
docID=0

for name in files:
    docID+=1
    buildIndex(name,docID)


#__Code to create the index file__

document = Document()
document.add_heading('Index', 0)

for element in dictG:
    document.add_paragraph(str(element),style='ListBullet')

#save index to a file
document.save('index.docx')

print(len(dictG))


#print('*****Inverted Index and Boolean retrival *********')
#Receive and Proccess the Query
processQuery()

